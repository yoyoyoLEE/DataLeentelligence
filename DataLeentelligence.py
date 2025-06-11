"""
Data Explorer with LLM Integration
Version: 1.0
Author: Dr. Yong Ha Lee
Description: Interactive data analysis tool with LLM integration for automated insights and abstract generation
Last Updated: June 2025
"""

import streamlit as st
import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
import seaborn as sns
import requests
import json
from io import BytesIO
from datetime import datetime
import os

from utilities import hash_password, verify_password, get_db_connection, log_activity, get_user_tier

# Login function
def login():
    st.title("Login")
    username = st.text_input("Username")
    password = st.text_input("Password", type="password")
    
    if st.button("Login"):
        conn = get_db_connection()
        try:
            cursor = conn.cursor()
            cursor.execute("SELECT id, password_hash, tier FROM users WHERE username = ?", (username,))
            user = cursor.fetchone()
            
            if user and verify_password(password, user[1]):
                st.session_state.logged_in = True
                st.session_state.username = username
                st.session_state.user_id = user[0]
                st.session_state.user_tier = user[2]
                
                # Update last login
                conn.execute("UPDATE users SET last_login = CURRENT_TIMESTAMP WHERE id = ?", (user[0],))
                conn.commit()
                
                # Log login activity
                log_activity(user[0], "login")
                
                st.rerun()
            else:
                st.error("Invalid username or password")
                log_activity(None, "failed_login", f"username: {username}")
        finally:
            conn.close()

# Check login state
if "logged_in" not in st.session_state:
    st.session_state.logged_in = False

if not st.session_state.logged_in:
    login()
    st.stop()

# Language settings
if 'language' not in st.session_state:
    st.session_state.language = 'english'

LANGUAGES = {
    'english': {
        'title': "Data Explorer with LLM",
        'upload': "Upload a CSV or Excel file",
        'success': "File uploaded successfully!",
        'stats': "Descriptive statistics",
        'abstract': "Generate Automatic Scientific Abstract",
        'clean': "Clean data (remove nulls, duplicates, basic normalization)",
        'edit': "Direct DataFrame editing",
        'advanced': "Advanced data operations",
        'question': "Ask a question about the data",
        'send': "Send request to LLM",
        'response': "LLM Response",
        'history': "LLM Response History",
        'clear': "Clear history",
        'export': "Export history CSV",
        'visuals': "View exploratory charts",
        'download': "Download modified file",
        'rows_to_analyze': "Number of rows to analyze"
    },
    'italian': {
        'title': "Data Explorer con LLM",
        'upload': "Carica un file CSV o Excel",
        'success': "File caricato con successo!",
        'stats': "Statistiche descrittive",
        'abstract': "Genera Abstract Scientifico Automatico",
        'clean': "Esegui pulizia dei dati (rimozione nulli, duplicati, normalizzazione base)",
        'edit': "Modifica diretta del DataFrame",
        'advanced': "Operazioni avanzate sui dati",
        'question': "Fai a domanda sui dati",
        'send': "Invia richiesta al LLM",
        'response': "Risposta dell'LLM",
        'history': "Cronologia delle Risposte dell'LLM",
        'clear': "Cancella cronologia",
        'export': "Esporta cronologia CSV",
        'visuals': "Visualizza grafici esplorativi",
        'download': "Scarica file modificato",
        'rows_to_analyze': "Numero di righe da analizzare"
    }
}

API_URL = "https://openrouter.ai/api/v1/chat/completions"
API_KEY = st.secrets["OPENROUTER_API_KEY"]

st.set_page_config(page_title="Data Explorer with LLM")

# Admin and language controls
col1, col2, col3 = st.columns([3,1,1])
with col1:
    st.title(LANGUAGES[st.session_state.language]['title'])
with col2:
    lang = st.selectbox("Language", options=['english', 'italian'], index=0 if st.session_state.language == 'english' else 1)
    if lang != st.session_state.language:
        st.session_state.language = lang
        st.rerun()
with col3:
    if st.session_state.user_tier == 'admin':
        if st.button("Admin Panel"):
            st.session_state.show_admin = not st.session_state.get('show_admin', False)
            st.rerun()

# Admin Panel
if st.session_state.get('show_admin', False) and st.session_state.user_tier == 'admin':
    st.sidebar.title("Admin Panel")
    
    admin_tab = st.sidebar.radio("Menu", ["User Management", "Activity Logs", "System Settings"])
    
    if admin_tab == "User Management":
        st.header("User Management")
        
        # Add new user
        with st.expander("Add New User"):
            new_username = st.text_input("Username")
            new_password = st.text_input("Password", type="password")
            new_tier = st.selectbox("Tier", ["admin", "tier2", "tier1"])
            if st.button("Add User"):
                conn = get_db_connection()
                try:
                    conn.execute(
                        "INSERT INTO users (username, password_hash, tier) VALUES (?, ?, ?)",
                        (new_username, hash_password(new_password), new_tier)
                    )
                    conn.commit()
                    st.success(f"User {new_username} added successfully")
                    log_activity(st.session_state.user_id, "user_add", f"added user: {new_username}")
                except sqlite3.IntegrityError:
                    st.error("Username already exists")
                finally:
                    conn.close()
        
        # User list and management
        st.subheader("Existing Users")
        conn = get_db_connection()
        users = conn.execute("SELECT id, username, tier, created_at, last_login FROM users").fetchall()
        conn.close()
        
        for user in users:
            with st.expander(f"{user[1]} ({user[2]})"):
                st.write(f"Created: {user[3]}")
                st.write(f"Last login: {user[4] or 'Never'}")
                
                new_tier = st.selectbox(
                    "Change Tier",
                    ["admin", "tier2", "tier1"],
                    index=["admin", "tier2", "tier1"].index(user[2]),
                    key=f"tier_{user[0]}"
                )
                
                if st.button("Update", key=f"update_{user[0]}"):
                    conn = get_db_connection()
                    conn.execute(
                        "UPDATE users SET tier = ? WHERE id = ?",
                        (new_tier, user[0])
                    )
                    conn.commit()
                    conn.close()
                    st.success("User updated")
                    log_activity(st.session_state.user_id, "user_update", f"updated user {user[1]} to tier {new_tier}")
                
                if st.button("Delete", key=f"delete_{user[0]}"):
                    conn = get_db_connection()
                    conn.execute("DELETE FROM users WHERE id = ?", (user[0],))
                    conn.commit()
                    conn.close()
                    st.success("User deleted")
                    log_activity(st.session_state.user_id, "user_delete", f"deleted user {user[1]}")
    
    elif admin_tab == "Activity Logs":
        st.header("Activity Logs")
        
        # Filter options
        col1, col2 = st.columns(2)
        with col1:
            user_filter = st.selectbox("Filter by user", ["All"] + [u[1] for u in users])
        with col2:
            action_filter = st.selectbox("Filter by action", ["All", "login", "failed_login", "user_add", "user_update", "user_delete"])
        
        # Query logs with filters
        conn = get_db_connection()
        query = "SELECT l.timestamp, u.username, l.action_type, l.details FROM activity_logs l LEFT JOIN users u ON l.user_id = u.id"
        params = []
        
        if user_filter != "All":
            query += " WHERE u.username = ?"
            params.append(user_filter)
            if action_filter != "All":
                query += " AND l.action_type = ?"
                params.append(action_filter)
        elif action_filter != "All":
            query += " WHERE l.action_type = ?"
            params.append(action_filter)
            
        query += " ORDER BY l.timestamp DESC LIMIT 100"
        
        logs = conn.execute(query, params).fetchall()
        conn.close()
        
        # Display logs
        for log in logs:
            st.write(f"{log[0]} | {log[1]} | {log[2]} | {log[3]}")
    
    elif admin_tab == "System Settings":
        st.header("System Settings")
        st.write("Configuration options coming soon")

SAVE_DIR = "./modifiche_auto_salvate"
os.makedirs(SAVE_DIR, exist_ok=True)

# Initialize user-specific chat history if not exists
if "user_chat_histories" not in st.session_state:
    st.session_state.user_chat_histories = {}

# Get or create chat history for current user
username = st.session_state.get('username', 'default')
if username not in st.session_state.user_chat_histories:
    st.session_state.user_chat_histories[username] = []

if "trigger_llm" not in st.session_state:
    st.session_state.trigger_llm = False

# Main tabs
if st.session_state.user_tier in ['admin', 'tier2']:
    tab1, tab2, tab3 = st.tabs(["Data Analysis", "AI Data Cleaning", "Data Retrieval"])
elif st.session_state.user_tier == 'tier1':
    tab1 = st.tabs(["Data Analysis"])[0]
else:
    tab1 = st.tabs(["Data Analysis"])[0]

# Data Retrieval Tab
if st.session_state.user_tier in ['admin', 'tier2'] and 'tab3' in locals():
    with tab3:
        st.header("Data Retrieval from Documents")
        
        # File upload
        uploaded_docs = st.file_uploader(
            "Upload PDF, JPG, or other documents", 
            type=["pdf", "jpg", "jpeg", "png", "docx"],
            accept_multiple_files=True
        )
        
        if uploaded_docs:
            st.success(f"{len(uploaded_docs)} document(s) uploaded")
            
            # Document processing
            with st.expander("Document Processing"):
                if st.button("Extract Data from Documents"):
                    st.session_state.extracted_data = []
                    
                    for doc in uploaded_docs:
                        # Placeholder for actual document processing
                        # In a real implementation, you would use:
                        # - PyPDF2 for PDFs
                        # - pytesseract for images
                        # - python-docx for Word docs
                        
                        extracted = {
                            "filename": doc.name,
                            "content": f"Extracted text from {doc.name} (placeholder)",
                            "tables": []
                        }
                        st.session_state.extracted_data.append(extracted)
                    
                    st.success("Data extraction complete")
                    log_activity(st.session_state.user_id, "data_extraction", f"Extracted data from {len(uploaded_docs)} documents")
            
            # Preview and confirmation
            if 'extracted_data' in st.session_state:
                st.subheader("Extracted Data Preview")
                
                for doc_data in st.session_state.extracted_data:
                    with st.expander(doc_data['filename']):
                        st.write(doc_data['content'])
                        
                        if doc_data['tables']:
                            st.write("Extracted Tables:")
                            for table in doc_data['tables']:
                                st.dataframe(table)
                
                # Data import options
                st.subheader("Import Options")
                import_action = st.radio(
                    "Select import action",
                    ["Create new dataset", "Append to current dataset"]
                )
                
                if st.button("Confirm Import"):
                    if import_action == "Create new dataset":
                        # Create new DataFrame from extracted data
                        st.session_state.df = pd.DataFrame()  # Replace with actual extracted data
                        st.success("New dataset created from documents")
                    else:
                        # Append to existing DataFrame
                        st.warning("Append functionality not yet implemented")
                    
                    log_activity(st.session_state.user_id, "data_import", f"Imported data from documents ({import_action})")

with tab1:
    uploaded_file = st.file_uploader(LANGUAGES[st.session_state.language]['upload'], type=["csv", "xlsx"])
    if uploaded_file is not None:
        try:
            if uploaded_file.name.endswith(".csv"):
                st.session_state.df = pd.read_csv(uploaded_file)
            else:
                st.session_state.df = pd.read_excel(uploaded_file)
            
            # Handle data types more carefully
            for col in st.session_state.df.columns:
                # Try to convert to numeric first
                try:
                    st.session_state.df[col] = pd.to_numeric(st.session_state.df[col], errors='ignore')
                except:
                    pass
                
                # Convert remaining non-numeric columns to string
                if not np.issubdtype(st.session_state.df[col].dtype, np.number):
                    st.session_state.df[col] = st.session_state.df[col].astype(str)
                    
            st.session_state.df = st.session_state.df.fillna('')
            st.success(LANGUAGES[st.session_state.language]['success'])
            st.dataframe(st.session_state.df.head())
        except Exception as e:
            st.error(f"Errore nel caricamento del file: {e}")
            st.stop()

if st.session_state.user_tier in ['admin', 'tier2'] and 'tab2' in locals():
    with tab2:
        if 'df' in st.session_state:
            st.header("AI-Assisted Data Cleaning")
            
            # AI Data Quality Assessment
            with st.expander("Data Quality Report"):
                if st.button("Generate Quality Report"):
                    if 'df' not in st.session_state or st.session_state.df.empty:
                        st.error("No data available for analysis")
                        st.stop()
                    
                    try:
                        sample_data = st.session_state.df.head(50).to_csv(index=False)
                        quality_prompt = f"""Analyze this dataset for quality issues:
{sample_data}

Provide a detailed report on:
1. Missing values per column
2. Data type inconsistencies
3. Potential outliers
4. Duplicate entries
5. Any other data quality concerns

Format as markdown with sections for each issue type."""
                    except Exception as e:
                        st.error(f"Error preparing data for analysis: {str(e)}")
                        st.stop()
                    
                    headers = {
                        "Authorization": f"Bearer {API_KEY}",
                        "Content-Type": "application/json"
                    }
                    payload = {
                        "model": "deepseek/deepseek-r1-0528-qwen3-8b:free",
                        "messages": [
                            {"role": "system", "content": "You are a data quality expert. Provide clear, actionable insights about data quality issues in the provided dataset."},
                            {"role": "user", "content": quality_prompt}
                        ],
                        "max_tokens": 2000,
                        "temperature": 0.3
                    }
                    
                    try:
                        response = requests.post(API_URL, headers=headers, json=payload)
                        response.raise_for_status()
                        result = response.json()
                        quality_report = result["choices"][0]["message"]["content"]
                        st.markdown(quality_report)
                        log_activity(st.session_state.user_id, "quality_report", "Generated data quality report")
                    except Exception as e:
                        st.error(f"Error generating quality report: {e}")
            
            # AI Suggested Cleaning
            with st.expander("Suggested Cleaning Operations"):
                if st.button("Get Cleaning Suggestions"):
                    cleaning_prompt = f"""Suggest cleaning operations for this dataset:
{st.session_state.df.head(50).to_csv(index=False)}

For each suggested operation, include:
1. Description of the issue
2. Recommended solution
3. Expected impact
4. Python/pandas code to implement (if applicable)

Format as markdown with clear sections."""
                    
                    payload["messages"] = [
                        {"role": "system", "content": "You are a data cleaning expert. Provide clear, actionable cleaning recommendations for the provided dataset."},
                        {"role": "user", "content": cleaning_prompt}
                    ]
                    
                    try:
                        response = requests.post(API_URL, headers=headers, json=payload)
                        response.raise_for_status()
                        result = response.json()
                        cleaning_suggestions = result["choices"][0]["message"]["content"]
                        st.markdown(cleaning_suggestions)
                        log_activity(st.session_state.user_id, "cleaning_suggestions", "Generated cleaning suggestions")
                    except Exception as e:
                        st.error(f"Error generating cleaning suggestions: {e}")
            
            # Interactive Cleaning
            with st.expander("Interactive Cleaning"):
                st.write("Perform specific cleaning operations:")
                
                if st.button("Auto-clean Common Issues"):
                    # Basic cleaning operations
                    df_clean = st.session_state.df.copy()
                    df_clean = df_clean.drop_duplicates()
                    df_clean = df_clean.dropna(how='all')
                    for col in df_clean.columns:
                        if df_clean[col].dtype == 'object':
                            df_clean[col] = df_clean[col].str.strip()
                    
                    st.session_state.df_clean_preview = df_clean
                    st.success("Basic cleaning applied. Preview below:")
                    st.dataframe(df_clean.head())
                    log_activity(st.session_state.user_id, "basic_cleaning", "Applied basic cleaning operations")
                
                if 'df_clean_preview' in st.session_state:
                    st.subheader("Cleaning Preview")
                    st.dataframe(st.session_state.df_clean_preview.head())
                    
                    if st.button("Apply Cleaning"):
                        st.session_state.df = st.session_state.df_clean_preview
                        st.success("Cleaning applied to dataset!")
                        log_activity(st.session_state.user_id, "cleaning_applied", "Applied cleaning operations")

if uploaded_file is not None and not st.session_state.get('show_admin', False):
    # Tier-based feature access
    if st.session_state.user_tier == 'tier1':
        # Limit tier1 to basic features
        if st.checkbox(LANGUAGES[st.session_state.language]['clean']):
            st.session_state.df = st.session_state.df.drop_duplicates()
            st.session_state.df = st.session_state.df.dropna(axis=0, how='all')
            st.session_state.df = st.session_state.df.apply(lambda x: x.str.strip() if x.dtype == 'object' else x.astype(str).str.strip())
            st.success("Pulizia completata")
            st.dataframe(st.session_state.df.head())

        if st.checkbox(LANGUAGES[st.session_state.language]['edit']):
            st.markdown("Manual editing (use Excel for large volumes)" if st.session_state.language == 'english' else "Modifica manuale (usa Excel per grandi volumi)")
            edited_df = st.experimental_data_editor(st.session_state.df, num_rows="dynamic")
            st.session_state.df = edited_df
            st.success("Modifiche applicate")
            st.dataframe(st.session_state.df.head())

    else:  # admin or tier2
        try:
            if uploaded_file.name.endswith(".csv"):
                st.session_state.df = pd.read_csv(uploaded_file)
            else:
                st.session_state.df = pd.read_excel(uploaded_file)
            
            # Handle data types more carefully
            for col in st.session_state.df.columns:
                # Try to convert to numeric first
                try:
                    st.session_state.df[col] = pd.to_numeric(st.session_state.df[col], errors='ignore')
                except:
                    pass
                
                # Convert remaining non-numeric columns to string
                if not np.issubdtype(st.session_state.df[col].dtype, np.number):
                    st.session_state.df[col] = st.session_state.df[col].astype(str)
                    
            st.session_state.df = st.session_state.df.fillna('')
        except Exception as e:
            st.error(f"Errore nel caricamento del file: {e}")
            st.stop()

    st.success(LANGUAGES[st.session_state.language]['success'])
    # Show file preview
    st.dataframe(st.session_state.df.head())
    
    st.subheader(LANGUAGES[st.session_state.language]['stats'])
    # Convert all columns to string for safe display
    display_df = st.session_state.df.astype(str)
    st.dataframe(display_df.describe(include='all'))

    rows_to_analyze = st.slider(
        LANGUAGES[st.session_state.language]['rows_to_analyze'],
        min_value=10,
        max_value=1000,
        value=100,
        step=10
    )

    if st.button(LANGUAGES[st.session_state.language]['abstract']):
        if st.session_state.language == 'english':
            abstract_prompt = f"""Please generate a complete scientific abstract structured with: Introduction, Methods, Results and Conclusion, based on:
1. This dataset (first {rows_to_analyze} rows):
{st.session_state.df.head(rows_to_analyze).to_csv(index=False)}

2. The user's question: {user_question if 'user_question' in locals() else 'No specific question provided'}

3. The AI's analysis: {st.session_state.latest_answer if 'latest_answer' in st.session_state else 'No previous analysis available'}

Focus on synthesizing these three elements into a coherent abstract."""
        else:
            abstract_prompt = f"""Genera un abstract scientifico completo strutturato in: Introduzione, Metodi, Risultati e Conclusione, basato su:
1. Questo dataset (prime {rows_to_analyze} righe):
{st.session_state.df.head(rows_to_analyze).to_csv(index=False)}

2. La domanda dell'utente: {user_question if 'user_question' in locals() else 'Nessuna domanda specifica'}

3. L'analisi precedente dell'AI: {st.session_state.latest_answer if 'latest_answer' in st.session_state else 'Nessuna analisi precedente'}

Concentrati sulla sintesi di questi tre elementi in un abstract coerente."""
        headers = {
            "Authorization": f"Bearer {API_KEY}",
            "Content-Type": "application/json"
        }
        payload = {
            "model": "deepseek/deepseek-r1-0528-qwen3-8b:free",
            "messages": [
                {"role": "system", "content": "You are a scientific assistant expert in writing structured abstracts and articles for academic journals (IMRaD format). Write in proper English with clear structure: Introduction, Methods, Results and Discussion. Use precise scientific language." if st.session_state.language == 'english' else 
                "Sei un assistente scientifico esperte nella redazione di abstract e articoli strutturati per riviste accademiche (formato IMRaD). Scrivi in italiano corretto e formale, con struttura chiara: Introduzione, Metodi, Risultati e Discussione. Evita traduzioni letterali dall'inglese e mantieni un linguaggio scientifico preciso."},
                {"role": "user", "content": abstract_prompt}
            ],
            "max_tokens": 4000,
            "temperature": 0.7
        }
        try:
            response = requests.post(API_URL, headers=headers, json=payload)
            response.raise_for_status()
            result = response.json()
            abstract = result["choices"][0]["message"]["content"]
            st.text_area("Structured Scientific Abstract" if st.session_state.language == 'english' else "Abstract Scientifico Strutturato", 
                        value=abstract, height=400)
            st.download_button("Download Abstract (TXT)" if st.session_state.language == 'english' else "Scarica Abstract (TXT)", 
                             abstract.encode("utf-8"), 
                             file_name="scientific_abstract.txt" if st.session_state.language == 'english' else "abstract_scientifico.txt", 
                             mime="text/plain")
        except Exception as e:
            st.error(f"Errore nella generazione dell'abstract: {e}")
    if st.checkbox(LANGUAGES[st.session_state.language]['clean']):
        st.session_state.df = st.session_state.df.drop_duplicates()
        st.session_state.df = st.session_state.df.dropna(axis=0, how='all')
        # Ensure all columns are treated as strings
        st.session_state.df = st.session_state.df.apply(lambda x: x.str.strip() if x.dtype == 'object' else x.astype(str).str.strip())
        st.success("Pulizia completata")
        st.dataframe(st.session_state.df.head())

    if st.checkbox(LANGUAGES[st.session_state.language]['edit']):
        st.markdown("Manual editing (use Excel for large volumes)" if st.session_state.language == 'english' else "Modifica manuale (usa Excel per grandi volumi)")
        edited_df = st.experimental_data_editor(st.session_state.df, num_rows="dynamic")
        st.session_state.df = edited_df
        st.success("Modifiche applicate")
        st.dataframe(st.session_state.df.head())

        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        auto_path = os.path.join(SAVE_DIR, f"autosave_{timestamp}.csv")
        st.session_state.df.to_csv(auto_path, index=False)
        st.info(f"Modifiche salvate automaticamente in: {auto_path}")

    if st.checkbox(LANGUAGES[st.session_state.language]['advanced']):
        if st.button("Group by column" if st.session_state.language == 'english' else "Raggruppa per colonna"):
            group_col = st.selectbox("Select column to group by" if st.session_state.language == 'english' else "Seleziona colonna per raggruppamento", st.session_state.df.columns)
            st.dataframe(st.session_state.df.groupby(group_col).size().reset_index(name='Count' if st.session_state.language == 'english' else 'Conteggio'))
        if st.button("Filter rows" if st.session_state.language == 'english' else "Filtra righe"):
            filter_col = st.selectbox("Select column to filter" if st.session_state.language == 'english' else "Seleziona colonna per filtrare", st.session_state.df.columns)
            filter_val = st.text_input("Value to search" if st.session_state.language == 'english' else "Valore da cercare")
            if filter_val:
                st.dataframe(st.session_state.df[st.session_state.df[filter_col].astype(str).str.contains(filter_val, case=False)])
        if st.button("Merge two columns" if st.session_state.language == 'english' else "Unisci due colonne"):
            col1 = st.selectbox("Prima colonna", st.session_state.df.columns, key='col1')
            col2 = st.selectbox("Seconda colonna", st.session_state.df.columns, key='col2')
            sep = st.text_input("Separatore", value="_")
            new_col_name = st.text_input("Nome nuova colonna", value=f"{col1}_{col2}")
            if col1 and col2 and new_col_name:
                st.session_state.df[new_col_name] = st.session_state.df[col1].astype(str) + sep + st.session_state.df[col2].astype(str)
                st.success(f"Creata nuova colonna: {new_col_name}")
                st.dataframe(st.session_state.df.head())

    user_question = st.text_area(LANGUAGES[st.session_state.language]['question'])

    if st.button(LANGUAGES[st.session_state.language]['send']) and user_question:
        context = f"Ecco i primi {rows_to_analyze} record del dataset:\n{st.session_state.df.head(rows_to_analyze).to_csv(index=False)}\n\nDomanda: {user_question}"

        headers = {
            "Authorization": f"Bearer {API_KEY}",
            "Content-Type": "application/json"
        }

        # Enhanced system prompt for statistical analysis
        system_prompt = """You are an AI assistant expert in comprehensive data analysis and scientific writing. 
        When statistical analysis is requested:
        1. Identify and perform all appropriate statistical tests including:
           - Descriptive statistics
           - T-tests (independent, paired)
           - ANOVA and MANOVA
           - Correlation and regression analysis
           - Non-parametric tests
           - Time series analysis
           - Multivariate analysis
        2. Explain the results and their interpretation
        3. Include assumptions checking and post-hoc tests when appropriate
        
        IMPORTANT: 
        - Do NOT include any Python code in your response
        - Only provide the analysis results and explanations
        - Use clear, non-technical language suitable for non-programmers
        
        Always respond in proper English with clear structure. Use formal but understandable language.""" if st.session_state.language == 'english' else """Sei un assistente AI esperto in analisi dati completa e redazione scientifica.
        Quando viene richiesta un'analisi statistica:
        1. Identifica ed esegui tutti i test statistici appropriati inclusi:
           - Statistiche descrittive
           - Test T (indipendenti, appaiati)
           - ANOVA e MANOVA
           - Analisi di correlazione e regressione
           - Test non parametrici
           - Analisi di serie temporali
           - Analisi multivariata
        2. Spiega i risultati e la loro interpretazione
        3. Includi verifica delle assunzioni e test post-hoc quando appropriato
        
        IMPORTANTE:
        - NON includere codice Python nella risposta
        - Fornisci solo i risultati dell'analisi e le spiegazioni
        - Usa un linguaggio chiaro e non tecnico adatto a non programmatori
        
        Scrivi sempre in italiano grammaticalmente perfetto, con struttura chiara."""
        
        payload = {
            "model": "deepseek/deepseek-r1-0528-qwen3-8b:free",
            "messages": [
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": context}
            ],
            "max_tokens": 3000,
            "temperature": 0.7  # More focused responses
        }

        try:
            response = requests.post(API_URL, headers=headers, json=payload)
            response.raise_for_status()
            result = response.json()
            answer = result["choices"][0]["message"]["content"]
            st.session_state.user_chat_histories[username].append((user_question, answer))
            st.session_state.latest_answer = answer
        except Exception as e:
            st.session_state.latest_answer = f"Errore durante la chiamata all'API: {e}"

        if "latest_answer" in st.session_state:
            st.markdown(f"### {LANGUAGES[st.session_state.language]['response']}")
            # Remove all code blocks and technical implementation details
            import re
            cleaned_response = re.sub(r'```.*?```', '', st.session_state.latest_answer, flags=re.DOTALL)  # Remove all code blocks
            cleaned_response = re.sub(r'(Here(.*?)code(.*?):|### (Python|Codice) (code|Python|per l\'analisi))', '', cleaned_response, flags=re.IGNORECASE)  # Remove code explanations and headers
            cleaned_response = re.sub(r'(python|codice)(.*?)\n', '', cleaned_response, flags=re.IGNORECASE)  # Remove language mentions
            cleaned_response = re.sub(r'#+\s*(Analysis|Analisi)\s*code\s*#+', '', cleaned_response, flags=re.IGNORECASE)  # Remove analysis code headers
            cleaned_response = cleaned_response.strip()
            st.text_area(LANGUAGES[st.session_state.language]['response'], value=cleaned_response, height=200)

        col1, col2 = st.columns(2)
        with col1:
            if st.download_button("Esporta come TXT", st.session_state.latest_answer.encode("utf-8"), file_name="risposta_ai.txt", mime="text/plain"):
                pass
        with col2:
            from docx import Document
            doc = Document()
            doc.add_heading("Risposta AI", level=1)
            doc.add_paragraph(st.session_state.latest_answer)
            doc_buffer = BytesIO()
            doc.save(doc_buffer)
            doc_buffer.seek(0)
            st.download_button("Esporta come DOCX", doc_buffer, file_name="risposta_ai.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

        if st.session_state.user_chat_histories[username]:
            st.markdown(f"### {LANGUAGES[st.session_state.language]['history']}")
            for i, (q, a) in enumerate(reversed(st.session_state.user_chat_histories[username])):
                st.markdown(f"**Q{i+1}:** {q}")
                st.text_area(f"Risposta {i+1}", value=a, height=200, key=f"ans_{i}")

            col1, col2 = st.columns([1, 1])
            with col1:
                if st.button(LANGUAGES[st.session_state.language]['clear']):
                    st.session_state.user_chat_histories[username] = []
                    st.session_state.latest_answer = ""
                    st.rerun()
            with col2:
                if st.button(LANGUAGES[st.session_state.language]['export']):
                    buffer = BytesIO()
                    pd.DataFrame(st.session_state.user_chat_histories[username], columns=["Domanda", "Risposta"]).to_csv(buffer, index=False)
                    st.download_button("Download CSV", buffer.getvalue(), file_name="cronologia_chat.csv", mime="text/csv")

        if st.checkbox(LANGUAGES[st.session_state.language]['visuals']):
            numeric_cols = st.session_state.df.select_dtypes(include='number').columns.tolist()
            if numeric_cols:
                col1 = st.selectbox("Scegli una variabile numerica per l'istogramma", numeric_cols)
                fig, ax = plt.subplots()
                sns.histplot(st.session_state.df[col1], bins=30, kde=True, ax=ax)
                st.pyplot(fig)
            else:
                st.warning("Non ci sono colonne numeriche per il grafico")

        if st.button(LANGUAGES[st.session_state.language]['download']):
            buffer = BytesIO()
            if uploaded_file.name.endswith(".csv"):
                st.session_state.df.to_csv(buffer, index=False)
                st.download_button("Download CSV", buffer.getvalue(), file_name="dati_modificati.csv", mime="text/csv")
            else:
                with pd.ExcelWriter(buffer, engine='xlsxwriter') as writer:
                    st.session_state.df.to_excel(writer, index=False)
                buffer.seek(0)
                st.download_button("Download Excel", buffer.getvalue(), file_name="dati_modificati.xlsx", mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")

    # Footer
    st.markdown("---")
    st.markdown(
        '<div style="text-align: right; font-size: 0.8em; color: #666;">'
        'Made by Dr. Y. Lee'
        '</div>', 
        unsafe_allow_html=True
    )
